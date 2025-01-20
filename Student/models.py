from django.db import models
import random 
from datetime import datetime, timedelta
from django.core.exceptions import ValidationError 
from django.core.mail import send_mail
from django.utils.crypto import get_random_string
from django.core.validators import FileExtensionValidator
import os
from pptxtopdf import convert
import pythoncom
from rake_nltk import Rake
from PyPDF2 import PdfReader
from sumy.parsers.plaintext import PlaintextParser
from sumy.summarizers.lsa import LsaSummarizer
import docx
import nltk
from nltk.tokenize import PunktSentenceTokenizer
from sumy.nlp.tokenizers import Tokenizer
from django.conf import settings
from django.utils import timezone

"""
    Login New Student
"""
from django.core.exceptions import ValidationError
from django.core.mail import send_mail

# Create your models here.

class LoginForm(models.Model):
    name = models.CharField(null=False,max_length=50)
    email = models.EmailField(null=False,max_length=100,default=' ')
    user_password = models.CharField(null=False, max_length=8)
    confirm_password = models.CharField(null=False, max_length=8)
    is_active = models.BooleanField(default=True)
    user_otp = models.CharField(max_length=6, blank=True, null=True)
    generated_otp = models.CharField(max_length=6, blank=True, null=True)  
    otp_expiry = models.DateTimeField(blank=True, null=True)

    def __str__(self):
        return f'{self.name} ({self.email})'

    class Meta:
        verbose_name_plural = 'Sign-Up Student Account' 
        verbose_name_plural = 'Login Students'

        
    def generate_otp(self):
    
        self.generated_otp = str(random.randint(100000, 999999))
        self.otp_expiry = datetime.now() + timedelta(minutes=5)  
        self.save()

        self.send_otp_email()

    def send_otp_email(self):
        subject = 'Your OTP for Login'
        message = f'Hello {self.name},\n\nYour OTP for login is {self.generated_otp}. It will expire in 5 minutes.'
        from_email = 'jhss12ahariharan@gmail.com'
        
        send_mail(
            subject,
            message,
            from_email,
            [self.email],
            fail_silently=False,
        )
    def set_password(self, password):
        """Hashes and stores the password securely."""
        from django.contrib.auth.hashers import make_password
        self.user_password = make_password(password)
        self.save()

    def check_password(self, password):
        """Compares a plain text password with the stored hashed password."""
        from django.contrib.auth.hashers import check_password
        return check_password(password, self.user_password)

"""
    Getting Student Information 
"""
    

YEAR = [
    ('1', '1st Year'),
    ('2', '2nd Year'),
    ('3', '3rd Year'),
    ('4', '4th Year'),
    ('5', '5th Year'),
]

CHOOSE = [
    ('Beginner', 'Beginner'),
    ('Intermediate', 'Intermediate'),
    ('Advance', 'Advance'),
]

MARK_TYPE = [
    ('percentage', 'Percentage'),
    ('cgpa', 'CGPA'),
]

class Home(models.Model):
    student_name = models.OneToOneField(LoginForm,null=True,on_delete=models.CASCADE,related_name='Login_student')
    college_name = models.CharField(max_length=100)
    course = models.CharField(null=False, max_length=30)
    year = models.CharField(choices=YEAR, max_length=2) 
    CGPA = models.CharField(choices=MARK_TYPE, max_length=10)  
    student_choice = models.CharField(choices=CHOOSE, max_length=20) 
    
    cgpa_percentage = models.DecimalField(max_digits=5, decimal_places=2, blank=True, null=True)
    cgpa_number = models.DecimalField(max_digits=4, decimal_places=2, blank=True, null=True)

    def save(self, *args, **kwargs):
        if self.CGPA == 'percentage' and self.cgpa_percentage is None:
            raise ValueError("Percentage CGPA is required when CGPA type is 'percentage'")
        elif self.CGPA == 'cgpa' and self.cgpa_number is None:
            raise ValueError("Numeric CGPA is required when CGPA type is 'cgpa'")
        
        if self.CGPA == 'percentage' and self.cgpa_number is not None:
            raise ValueError("Please provide only CGPA percentage, not CGPA number.")
        elif self.CGPA == 'cgpa' and self.cgpa_percentage is not None:
            raise ValueError("Please provide only CGPA number, not CGPA percentage.")
        
        super(Home, self).save(*args, **kwargs)

    def __str__(self):
        return f"{self.college_name} - {self.course} - {self.year} - {self.CGPA}"
    
    class Meta:
        verbose_name_plural = 'Student Information'

""" 
    Sign-In Existing Student Account
"""
class StudentID(models.Model):
    student = models.OneToOneField(LoginForm,on_delete=models.CASCADE,related_name='student_name')
    unique_id = models.CharField(max_length=4,unique=True,editable=False)
    password = models.CharField(max_length=8)

    def save(self,*args,**kwargs):
        if not self.unique_id:
            while True:
                random_id = str(random.randint(1000, 9999))
                if not StudentID.objects.filter(unique_id=random_id).exists():
                    self.unique_id = random_id
                    break
        super().save(*args, **kwargs)

    @staticmethod
    def generate_unique_student_id():
        while True:
            random_id = get_random_string(length=4, allowed_chars='0123456789')
            if not StudentID.objects.filter(student_id=random_id).exists():
                return random_id
            
    def __str__(self):
        return f"StudentID: {self.unique_id} - {self.student.name}"
            
    class Meta:
        verbose_name_plural = 'Sign-In Student Account'


"""
    Uploading the File
"""
STUDENT_CHOICE = [
    ('Summary','Summary'),
    ('Keywords','Keywords'),
]

class UploadFile(models.Model):
    upload_file = models.FileField(null=False,upload_to='documents/',validators=[FileExtensionValidator(['pdf', 'docx','pptx'])])
    uploaded_at = models.DateTimeField(auto_now_add=True)
    student_options = models.CharField(choices=STUDENT_CHOICE,max_length=30)
    keywords = models.TextField(null=True, blank=True) 
    summary = models.TextField(null=True, blank=True)  

    class Meta:
        verbose_name_plural = 'File Upload'


    def extract_keywords_from_text(self, text):
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')
        r = Rake()
        r.extract_keywords_from_text(text)
        ranked_phrases = r.get_ranked_phrases_with_scores()
        return [phrase for score, phrase in ranked_phrases[:7]]

    def generate_summary_from_text(self, text, num_sentences=3):
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')
        tokenizer = PunktSentenceTokenizer()
        parser = PlaintextParser.from_string(text,  Tokenizer("english"))
        summarizer = LsaSummarizer()
        summary = summarizer(parser.document, num_sentences)
        return ' '.join([str(sentence) for sentence in summary])

    def convert_pptx_to_pdf(self):
        pptx_path = os.path.join(settings.MEDIA_ROOT, 'documents', os.path.basename(self.upload_file.name))
        pdf_dir = os.path.dirname(pptx_path)
        pdf_path = None
        try:
            pythoncom.CoInitialize()
            print(f"File path for conversion: {self.upload_file.path}")

            convert(pptx_path, pdf_dir)
            pdf_path = pptx_path.replace('.pptx', '.pdf')
            if os.path.exists(pdf_path):
                self.upload_file.name = self.upload_file.name.replace('.pptx', '.pdf')
                self.upload_file.save()
        except Exception as e:
            print(f"Error converting .pptx to .pdf: {e}")
        finally:
            pythoncom.CoUninitialize()

    def process_file(self):
        file_path = self.upload_file.path
        print(f"Processing file: {file_path}")  
        if self.student_options == 'Keywords':
            if file_path.endswith('.pdf'):
                self.keywords = self.extract_keywords_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                self.keywords = self.extract_keywords_from_docx(file_path)
            self.__class__.objects.filter(id=self.id).update(keywords=self.keywords)
        elif self.student_options == 'Summary':
            if file_path.endswith('.pdf'):
                self.summary = self.generate_summary_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                self.summary = self.generate_summary_from_docx(file_path)
            self.__class__.objects.filter(id=self.id).update(summary=self.summary) 

    def save(self, *args, **kwargs):
        
        if self.upload_file.name.endswith('.pptx') and not self.upload_file.name.endswith('.pdf'):
            self.convert_pptx_to_pdf()
        super().save(*args, **kwargs)  
        self.process_file()

    def extract_keywords_from_pdf(self, file_path):
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        return self.extract_keywords_from_text(text)

    def generate_summary_from_pdf(self, file_path, num_sentences=3):
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        return self.generate_summary_from_text(text, num_sentences)

    def extract_keywords_from_docx(self, file_path):
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return self.extract_keywords_from_text(text)

    def generate_summary_from_docx(self, file_path, num_sentences=3):
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return self.generate_summary_from_text(text, num_sentences)

class Feedback(models.Model):
    FEEDBACK_CHOICES = [
        ('summary', 'Summary'),
        ('keywords', 'Keywords'),
    ]

    text = models.TextField() 
    category = models.CharField(max_length=20, choices=FEEDBACK_CHOICES) 
    created_at = models.DateTimeField(auto_now_add=True)
    user = models.ForeignKey(LoginForm, on_delete=models.CASCADE, related_name="feedbacks")  
    
    def __str__(self):
        return f"Feedback for {self.category} by {self.user.name}"

class UserActivity(models.Model):
    ACTION_CHOICES = [
        ('login', 'Login'),
        ('logout', 'Logout'),
        ('file_upload', 'File Upload'),
        ('file_view', 'File View'),
        ('file_translate', 'File Translate'),
        ('file_speak', 'File Speak'),
        ('feedback', 'Feedback'),
        ('otp_verified', 'OTP Verified'),
    ]
    
    user = models.ForeignKey(Home, on_delete=models.CASCADE, related_name="user_activities")
    action = models.CharField(max_length=50, choices=ACTION_CHOICES)
    timestamp = models.DateTimeField(default=timezone.now)
    details = models.TextField(null=True, blank=True)  

    def __str__(self):
        return f"{self.user.name} - {self.get_action_display()} at {self.timestamp}"

    class Meta:
        verbose_name_plural = 'User Activities'